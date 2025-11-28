import sys
import os
import subprocess
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QPushButton, QLabel, QFileDialog, QComboBox, QMessageBox, 
    QDialog, QTabWidget, QGroupBox, QGridLayout, QLineEdit,
    QSpinBox, QFontComboBox, QCheckBox, QColorDialog, QFormLayout,
    QTextEdit, QSplitter, QListWidget, QStackedWidget, QButtonGroup,
    QRadioButton, QScrollArea, QMenuBar, QMenu, QDoubleSpinBox
)
from PyQt5.QtCore import Qt, pyqtSignal
from PyQt5.QtGui import QFont
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


class FormatConfigDialog(QDialog):
    """排版配置对话框"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("排版配置")
        self.setMinimumSize(800, 600)
        self.config = {}
        
        self.init_ui()
        self.load_default_config()
        
    def init_ui(self):
        main_layout = QVBoxLayout(self)  # 改为垂直布局，以便将预设模板区域放在顶部
        
        # 预设模板区域 - 移到顶部
        template_container = QWidget()
        template_layout = QVBoxLayout(template_container)
        
        # 预设模板标题和选择区域
        template_select_layout = QHBoxLayout()
        template_select_layout.addWidget(QLabel("预设模板："))
        
        self.template_combo = QComboBox()
        self.template_combo.addItems([
            "默认模板", "学术论文模板", "报告模板", "小说模板", "简历模板", "信函模板"
        ])
        self.template_combo.setCurrentText("默认模板")
        self.template_combo.currentTextChanged.connect(self.load_template_preview)
        template_select_layout.addWidget(self.template_combo)
        
        apply_template_button = QPushButton("应用模板")
        apply_template_button.clicked.connect(self.apply_preset_template)
        template_select_layout.addWidget(apply_template_button)
        
        # 导出模板按钮也放在这里
        export_template_button = QPushButton("导出模板")
        export_template_button.clicked.connect(self.export_template)
        template_select_layout.addWidget(export_template_button)
        
        template_layout.addLayout(template_select_layout)
        
        # 模板预览文本
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setMaximumHeight(120)
        self.preview_text.setPlainText("默认模板预览")
        template_layout.addWidget(self.preview_text)
        
        # 添加模板容器到主布局
        main_layout.addWidget(template_container)
        
        # 创建导航列表和配置区域的水平布局
        nav_config_layout = QHBoxLayout()
        
        # 创建导航列表和配置区域的分割器
        splitter = QSplitter(Qt.Horizontal)
        
        # 创建导航列表
        self.nav_list = QListWidget()
        self.nav_list.setMaximumWidth(200)
        
        # 添加导航项
        nav_items = [
            "基础文本与段落样式",
            "标题与层级样式", 
            "列表样式",
            "引用与交互元素样式",
            "非文本与布局样式块",
            "页面设置"
        ]
        
        for item in nav_items:
            self.nav_list.addItem(item)
        
        # 连接导航点击事件
        self.nav_list.currentRowChanged.connect(self.show_config_section)
        
        # 创建配置区域
        self.config_stack = QStackedWidget()
        
        # 创建各配置区域
        self.basic_text_section = self.create_basic_text_section()
        self.heading_section = self.create_heading_section()
        self.list_section = self.create_list_section()
        self.reference_section = self.create_reference_section()
        self.layout_section = self.create_layout_section()
        self.page_section = self.create_page_section()
        
        # 添加配置区域到堆栈
        self.config_stack.addWidget(self.basic_text_section)
        self.config_stack.addWidget(self.heading_section)
        self.config_stack.addWidget(self.list_section)
        self.config_stack.addWidget(self.reference_section)
        self.config_stack.addWidget(self.layout_section)
        self.config_stack.addWidget(self.page_section)
        
        # 设置默认选中第一项
        self.nav_list.setCurrentRow(0)
        
        # 添加到分割器
        splitter.addWidget(self.nav_list)
        splitter.addWidget(self.config_stack)
        splitter.setStretchFactor(0, 0)
        splitter.setStretchFactor(1, 1)
        
        # 添加分割器到导航配置布局
        nav_config_layout.addWidget(splitter)
        
        # 添加导航配置布局到主布局
        main_layout.addLayout(nav_config_layout)
        
        # 按钮区域
        button_layout = QHBoxLayout()
        cancel_button = QPushButton("取消")
        cancel_button.clicked.connect(self.reject)
        
        button_layout.addStretch()
        button_layout.addWidget(cancel_button)
        button_layout.addStretch()
        
        # 添加按钮布局到主布局
        main_layout.addLayout(button_layout)
    
    def show_config_section(self, index):
        """根据选择的导航项显示对应的配置区域"""
        self.config_stack.setCurrentIndex(index)
    
    def create_basic_text_section(self):
        """创建基础文本与段落样式配置区域"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 创建滚动区域
        scroll = QScrollArea()
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        
        # 段落样式
        paragraph_group = QGroupBox("段落")
        paragraph_layout = QFormLayout()
        
        self.paragraph_font_combo = QFontComboBox()
        self.paragraph_font_combo.setFontFilters(QFontComboBox.AllFonts)
        paragraph_layout.addRow("字体:", self.paragraph_font_combo)
        
        self.paragraph_font_size_spin = QSpinBox()
        self.paragraph_font_size_spin.setRange(8, 72)
        self.paragraph_font_size_spin.setValue(12)
        paragraph_layout.addRow("字体大小:", self.paragraph_font_size_spin)
        
        self.paragraph_bold_check = QCheckBox("粗体")
        paragraph_layout.addRow("", self.paragraph_bold_check)
        
        self.paragraph_italic_check = QCheckBox("斜体")
        paragraph_layout.addRow("", self.paragraph_italic_check)
        
        paragraph_group.setLayout(paragraph_layout)
        scroll_layout.addWidget(paragraph_group)
        
        # 正文样式
        body_group = QGroupBox("正文")
        body_layout = QFormLayout()
        
        self.body_font_combo = QFontComboBox()
        self.body_font_combo.setFontFilters(QFontComboBox.AllFonts)
        body_layout.addRow("字体:", self.body_font_combo)
        
        self.body_font_size_spin = QSpinBox()
        self.body_font_size_spin.setRange(8, 72)
        self.body_font_size_spin.setValue(12)
        body_layout.addRow("字体大小:", self.body_font_size_spin)
        
        self.body_bold_check = QCheckBox("粗体")
        body_layout.addRow("", self.body_bold_check)
        
        self.body_italic_check = QCheckBox("斜体")
        body_layout.addRow("", self.body_italic_check)
        
        body_group.setLayout(body_layout)
        scroll_layout.addWidget(body_group)
        
        # 正文字符样式
        char_group = QGroupBox("正文字符（包含中文、英文、数字的字体）")
        char_layout = QFormLayout()
        
        # 中文字体
        self.chinese_font_combo = QFontComboBox()
        self.chinese_font_combo.setFontFilters(QFontComboBox.AllFonts)
        char_layout.addRow("中文字体:", self.chinese_font_combo)
        
        # 英文字体
        self.english_font_combo = QFontComboBox()
        self.english_font_combo.setFontFilters(QFontComboBox.AllFonts)
        char_layout.addRow("英文字体:", self.english_font_combo)
        
        # 数字字体
        self.number_font_combo = QFontComboBox()
        self.number_font_combo.setFontFilters(QFontComboBox.AllFonts)
        char_layout.addRow("数字字体:", self.number_font_combo)
        
        self.char_font_size_spin = QSpinBox()
        self.char_font_size_spin.setRange(8, 72)
        self.char_font_size_spin.setValue(12)
        char_layout.addRow("字体大小:", self.char_font_size_spin)
        
        char_group.setLayout(char_layout)
        scroll_layout.addWidget(char_group)
        
        # 段落格式
        format_group = QGroupBox("段落格式")
        format_layout = QFormLayout()
        
        self.first_line_indent_spin = QDoubleSpinBox()
        self.first_line_indent_spin.setRange(0, 10)
        self.first_line_indent_spin.setValue(2)
        self.first_line_indent_spin.setSuffix(" 字符")
        format_layout.addRow("首行缩进:", self.first_line_indent_spin)
        
        self.line_spacing_combo = QComboBox()
        self.line_spacing_combo.addItems(["1.0", "1.15", "1.5", "2.0"])
        self.line_spacing_combo.setCurrentText("1.5")
        format_layout.addRow("行间距:", self.line_spacing_combo)
        
        self.paragraph_spacing_spin = QDoubleSpinBox()
        self.paragraph_spacing_spin.setRange(0, 48)
        self.paragraph_spacing_spin.setValue(6)
        self.paragraph_spacing_spin.setSuffix(" 磅")
        format_layout.addRow("段前间距:", self.paragraph_spacing_spin)
        
        # 对齐方式
        self.align_group = QButtonGroup()
        self.left_radio = QRadioButton("左对齐")
        self.center_radio = QRadioButton("居中对齐")
        self.right_radio = QRadioButton("右对齐")
        self.justify_radio = QRadioButton("两端对齐")
        self.justify_radio.setChecked(True)
        
        self.align_group.addButton(self.left_radio, 0)
        self.align_group.addButton(self.center_radio, 1)
        self.align_group.addButton(self.right_radio, 2)
        self.align_group.addButton(self.justify_radio, 3)
        
        align_layout = QHBoxLayout()
        align_layout.addWidget(self.left_radio)
        align_layout.addWidget(self.center_radio)
        align_layout.addWidget(self.right_radio)
        align_layout.addWidget(self.justify_radio)
        format_layout.addRow("对齐方式:", align_layout)
        
        format_group.setLayout(format_layout)
        scroll_layout.addWidget(format_group)
        
        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)
        
        return widget
    
    def create_heading_section(self):
        """创建标题与层级样式配置区域"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 创建滚动区域
        scroll = QScrollArea()
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        
        # 标题样式
        self.heading_configs = {}
        
        for level in range(1, 10):
            group = QGroupBox(f"标题{level}")
            group_layout = QFormLayout()
            
            # 字体选择
            font_combo = QFontComboBox()
            font_combo.setFontFilters(QFontComboBox.AllFonts)
            group_layout.addRow("字体:", font_combo)
            
            # 字体大小
            size_spin = QSpinBox()
            size_spin.setRange(6, 72)
            if level == 1:
                size_spin.setValue(16)
            elif level == 2:
                size_spin.setValue(14)
            elif level == 3:
                size_spin.setValue(13)
            else:
                size_spin.setValue(12)
            group_layout.addRow("字体大小:", size_spin)
            
            # 粗体
            bold_check = QCheckBox("粗体")
            bold_check.setChecked(True if level <= 3 else False)
            group_layout.addRow("", bold_check)
            
            # 斜体
            italic_check = QCheckBox("斜体")
            group_layout.addRow("", italic_check)
            
            # 下划线
            underline_check = QCheckBox("下划线")
            group_layout.addRow("", underline_check)
            
            # 对齐方式
            align_combo = QComboBox()
            align_combo.addItems(["左对齐", "居中对齐", "右对齐", "两端对齐"])
            if level == 1:
                align_combo.setCurrentText("居中对齐")
            else:
                align_combo.setCurrentText("左对齐")
            group_layout.addRow("对齐方式:", align_combo)
            
            # 保存控件引用
            self.heading_configs[f"标题{level}"] = {
                'font': font_combo,
                'size': size_spin,
                'bold': bold_check,
                'italic': italic_check,
                'underline': underline_check,
                'align': align_combo
            }
            
            group.setLayout(group_layout)
            scroll_layout.addWidget(group)
        
        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)
        
        return widget
    
    def create_list_section(self):
        """创建列表样式配置区域"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 创建滚动区域
        scroll = QScrollArea()
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        
        # 有序列表样式
        ordered_group = QGroupBox("有序列表")
        ordered_layout = QFormLayout()
        
        self.ordered_font_combo = QFontComboBox()
        self.ordered_font_combo.setFontFilters(QFontComboBox.AllFonts)
        ordered_layout.addRow("字体:", self.ordered_font_combo)
        
        self.ordered_font_size_spin = QSpinBox()
        self.ordered_font_size_spin.setRange(8, 72)
        self.ordered_font_size_spin.setValue(12)
        ordered_layout.addRow("字体大小:", self.ordered_font_size_spin)
        
        self.ordered_bold_check = QCheckBox("粗体")
        ordered_layout.addRow("", self.ordered_bold_check)
        
        self.ordered_italic_check = QCheckBox("斜体")
        ordered_layout.addRow("", self.ordered_italic_check)
        
        # 编号格式
        self.ordered_format_combo = QComboBox()
        self.ordered_format_combo.addItems(["1, 2, 3...", "a, b, c...", "A, B, C...", "I, II, III..."])
        ordered_layout.addRow("编号格式:", self.ordered_format_combo)
        
        ordered_group.setLayout(ordered_layout)
        scroll_layout.addWidget(ordered_group)
        
        # 无序列表样式
        unordered_group = QGroupBox("无序列表")
        unordered_layout = QFormLayout()
        
        self.unordered_font_combo = QFontComboBox()
        self.unordered_font_combo.setFontFilters(QFontComboBox.AllFonts)
        unordered_layout.addRow("字体:", self.unordered_font_combo)
        
        self.unordered_font_size_spin = QSpinBox()
        self.unordered_font_size_spin.setRange(8, 72)
        self.unordered_font_size_spin.setValue(12)
        unordered_layout.addRow("字体大小:", self.unordered_font_size_spin)
        
        self.unordered_bold_check = QCheckBox("粗体")
        unordered_layout.addRow("", self.unordered_bold_check)
        
        self.unordered_italic_check = QCheckBox("斜体")
        unordered_layout.addRow("", self.unordered_italic_check)
        
        # 项目符号格式
        self.unordered_format_combo = QComboBox()
        self.unordered_format_combo.addItems(["●", "■", "○", "■", "►"])
        unordered_layout.addRow("项目符号:", self.unordered_format_combo)
        
        unordered_group.setLayout(unordered_layout)
        scroll_layout.addWidget(unordered_group)
        
        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)
        
        return widget
    
    def create_reference_section(self):
        """创建引用与交互元素样式配置区域"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 创建滚动区域
        scroll = QScrollArea()
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        
        # 超链接样式
        link_group = QGroupBox("超链接")
        link_layout = QFormLayout()
        
        self.link_font_combo = QFontComboBox()
        self.link_font_combo.setFontFilters(QFontComboBox.AllFonts)
        link_layout.addRow("字体:", self.link_font_combo)
        
        self.link_font_size_spin = QSpinBox()
        self.link_font_size_spin.setRange(8, 72)
        self.link_font_size_spin.setValue(12)
        link_layout.addRow("字体大小:", self.link_font_size_spin)
        
        self.link_bold_check = QCheckBox("粗体")
        link_layout.addRow("", self.link_bold_check)
        
        self.link_italic_check = QCheckBox("斜体")
        link_layout.addRow("", self.link_italic_check)
        
        self.link_underline_check = QCheckBox("下划线")
        self.link_underline_check.setChecked(True)
        link_layout.addRow("", self.link_underline_check)
        
        # 链接颜色
        self.link_color_button = QPushButton("选择颜色")
        self.link_color_button.clicked.connect(lambda: self.choose_color(self.link_color_button))
        link_layout.addRow("链接颜色:", self.link_color_button)
        
        link_group.setLayout(link_layout)
        scroll_layout.addWidget(link_group)
        
        # 脚注样式
        footnote_group = QGroupBox("脚注")
        footnote_layout = QFormLayout()
        
        self.footnote_font_combo = QFontComboBox()
        self.footnote_font_combo.setFontFilters(QFontComboBox.AllFonts)
        footnote_layout.addRow("字体:", self.footnote_font_combo)
        
        self.footnote_font_size_spin = QSpinBox()
        self.footnote_font_size_spin.setRange(6, 48)
        self.footnote_font_size_spin.setValue(9)
        footnote_layout.addRow("字体大小:", self.footnote_font_size_spin)
        
        self.footnote_bold_check = QCheckBox("粗体")
        footnote_layout.addRow("", self.footnote_bold_check)
        
        self.footnote_italic_check = QCheckBox("斜体")
        self.footnote_italic_check.setChecked(True)
        footnote_layout.addRow("", self.footnote_italic_check)
        
        self.footnote_underline_check = QCheckBox("下划线")
        footnote_layout.addRow("", self.footnote_underline_check)
        
        footnote_group.setLayout(footnote_layout)
        scroll_layout.addWidget(footnote_group)
        
        # 脚注引用文本样式
        footnote_ref_group = QGroupBox("脚注引用文本")
        footnote_ref_layout = QFormLayout()
        
        self.footnote_ref_font_size_spin = QSpinBox()
        self.footnote_ref_font_size_spin.setRange(6, 24)
        self.footnote_ref_font_size_spin.setValue(9)
        footnote_ref_layout.addRow("字体大小:", self.footnote_ref_font_size_spin)
        
        self.footnote_ref_superscript_check = QCheckBox("上标")
        self.footnote_ref_superscript_check.setChecked(True)
        footnote_ref_layout.addRow("", self.footnote_ref_superscript_check)
        
        self.footnote_ref_bold_check = QCheckBox("粗体")
        footnote_ref_layout.addRow("", self.footnote_ref_bold_check)
        
        self.footnote_ref_italic_check = QCheckBox("斜体")
        footnote_ref_layout.addRow("", self.footnote_ref_italic_check)
        
        footnote_ref_group.setLayout(footnote_ref_layout)
        scroll_layout.addWidget(footnote_ref_group)
        
        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)
        
        return widget
    
    def create_layout_section(self):
        """创建非文本与布局样式块配置区域"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 创建滚动区域
        scroll = QScrollArea()
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        
        # 块文本样式
        block_group = QGroupBox("块文本")
        block_layout = QFormLayout()
        
        self.block_font_combo = QFontComboBox()
        self.block_font_combo.setFontFilters(QFontComboBox.AllFonts)
        block_layout.addRow("字体:", self.block_font_combo)
        
        self.block_font_size_spin = QSpinBox()
        self.block_font_size_spin.setRange(8, 72)
        self.block_font_size_spin.setValue(12)
        block_layout.addRow("字体大小:", self.block_font_size_spin)
        
        self.block_bold_check = QCheckBox("粗体")
        block_layout.addRow("", self.block_bold_check)
        
        self.block_italic_check = QCheckBox("斜体")
        block_layout.addRow("", self.block_italic_check)
        
        # 块文本缩进
        self.block_indent_spin = QDoubleSpinBox()
        self.block_indent_spin.setRange(0, 10)
        self.block_indent_spin.setValue(2)
        self.block_indent_spin.setSuffix(" 字符")
        block_layout.addRow("左右缩进:", self.block_indent_spin)
        
        block_group.setLayout(block_layout)
        scroll_layout.addWidget(block_group)
        
        # 水平线样式
        hr_group = QGroupBox("水平线")
        hr_layout = QFormLayout()
        
        self.hr_thickness_spin = QDoubleSpinBox()
        self.hr_thickness_spin.setRange(0.5, 5)
        self.hr_thickness_spin.setValue(1)
        self.hr_thickness_spin.setSuffix(" 磅")
        hr_layout.addRow("线条粗细:", self.hr_thickness_spin)
        
        self.hr_width_spin = QSpinBox()
        self.hr_width_spin.setRange(10, 100)
        self.hr_width_spin.setValue(100)
        self.hr_width_spin.setSuffix("%")
        hr_layout.addRow("线条宽度:", self.hr_width_spin)
        
        self.hr_color_button = QPushButton("选择颜色")
        self.hr_color_button.clicked.connect(lambda: self.choose_color(self.hr_color_button))
        hr_layout.addRow("线条颜色:", self.hr_color_button)
        
        hr_group.setLayout(hr_layout)
        scroll_layout.addWidget(hr_group)
        
        # 表标题样式
        table_caption_group = QGroupBox("表标题")
        table_caption_layout = QFormLayout()
        
        self.table_caption_font_combo = QFontComboBox()
        self.table_caption_font_combo.setFontFilters(QFontComboBox.AllFonts)
        table_caption_layout.addRow("字体:", self.table_caption_font_combo)
        
        self.table_caption_font_size_spin = QSpinBox()
        self.table_caption_font_size_spin.setRange(8, 48)
        self.table_caption_font_size_spin.setValue(12)
        table_caption_layout.addRow("字体大小:", self.table_caption_font_size_spin)
        
        self.table_caption_bold_check = QCheckBox("粗体")
        self.table_caption_bold_check.setChecked(True)
        table_caption_layout.addRow("", self.table_caption_bold_check)
        
        self.table_caption_italic_check = QCheckBox("斜体")
        table_caption_layout.addRow("", self.table_caption_italic_check)
        
        # 表标题位置
        self.table_caption_position_combo = QComboBox()
        self.table_caption_position_combo.addItems(["表格上方", "表格下方"])
        table_caption_layout.addRow("位置:", self.table_caption_position_combo)
        
        table_caption_group.setLayout(table_caption_layout)
        scroll_layout.addWidget(table_caption_group)
        
        # 表格样式
        table_group = QGroupBox("表格")
        table_layout = QFormLayout()
        
        self.table_font_combo = QFontComboBox()
        self.table_font_combo.setFontFilters(QFontComboBox.AllFonts)
        table_layout.addRow("字体:", self.table_font_combo)
        
        self.table_font_size_spin = QSpinBox()
        self.table_font_size_spin.setRange(8, 48)
        self.table_font_size_spin.setValue(10)
        table_layout.addRow("字体大小:", self.table_font_size_spin)
        
        self.table_header_bold_check = QCheckBox("表头粗体")
        self.table_header_bold_check.setChecked(True)
        table_layout.addRow("", self.table_header_bold_check)
        
        # 表格边框
        self.table_border_check = QCheckBox("显示边框")
        self.table_border_check.setChecked(True)
        table_layout.addRow("", self.table_border_check)
        
        table_group.setLayout(table_layout)
        scroll_layout.addWidget(table_group)
        
        # 图片说明样式
        figure_caption_group = QGroupBox("图片说明")
        figure_caption_layout = QFormLayout()
        
        self.figure_caption_font_combo = QFontComboBox()
        self.figure_caption_font_combo.setFontFilters(QFontComboBox.AllFonts)
        figure_caption_layout.addRow("字体:", self.figure_caption_font_combo)
        
        self.figure_caption_font_size_spin = QSpinBox()
        self.figure_caption_font_size_spin.setRange(8, 48)
        self.figure_caption_font_size_spin.setValue(10)
        figure_caption_layout.addRow("字体大小:", self.figure_caption_font_size_spin)
        
        self.figure_caption_bold_check = QCheckBox("粗体")
        figure_caption_layout.addRow("", self.figure_caption_bold_check)
        
        self.figure_caption_italic_check = QCheckBox("斜体")
        self.figure_caption_italic_check.setChecked(True)
        figure_caption_layout.addRow("", self.figure_caption_italic_check)
        
        # 图片说明位置
        self.figure_caption_position_combo = QComboBox()
        self.figure_caption_position_combo.addItems(["图片下方", "图片上方"])
        self.figure_caption_position_combo.setCurrentText("图片下方")
        figure_caption_layout.addRow("位置:", self.figure_caption_position_combo)
        
        figure_caption_group.setLayout(figure_caption_layout)
        scroll_layout.addWidget(figure_caption_group)
        
        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)
        
        return widget
    
    def create_page_section(self):
        """创建页面设置配置区域"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 页面尺寸
        size_group = QGroupBox("页面尺寸")
        size_layout = QFormLayout()
        
        self.page_size_combo = QComboBox()
        self.page_size_combo.addItems(["A4", "A3", "A5", "B4", "B5", "Letter", "Legal"])
        self.page_size_combo.setCurrentText("A4")
        size_layout.addRow("纸张大小:", self.page_size_combo)
        
        self.orientation_group = QButtonGroup()
        self.portrait_radio = QRadioButton("纵向")
        self.landscape_radio = QRadioButton("横向")
        self.portrait_radio.setChecked(True)
        
        self.orientation_group.addButton(self.portrait_radio, 0)
        self.orientation_group.addButton(self.landscape_radio, 1)
        
        orientation_layout = QHBoxLayout()
        orientation_layout.addWidget(self.portrait_radio)
        orientation_layout.addWidget(self.landscape_radio)
        size_layout.addRow("页面方向:", orientation_layout)
        
        size_group.setLayout(size_layout)
        layout.addWidget(size_group)
        
        # 页边距
        margin_group = QGroupBox("页边距")
        margin_layout = QFormLayout()
        
        self.margin_top_spin = QDoubleSpinBox()
        self.margin_top_spin.setRange(0, 10)
        self.margin_top_spin.setValue(2.54)
        self.margin_top_spin.setSuffix(" 厘米")
        margin_layout.addRow("上边距:", self.margin_top_spin)
        
        self.margin_bottom_spin = QDoubleSpinBox()
        self.margin_bottom_spin.setRange(0, 10)
        self.margin_bottom_spin.setValue(2.54)
        self.margin_bottom_spin.setSuffix(" 厘米")
        margin_layout.addRow("下边距:", self.margin_bottom_spin)
        
        self.margin_left_spin = QDoubleSpinBox()
        self.margin_left_spin.setRange(0, 10)
        self.margin_left_spin.setValue(3.17)
        self.margin_left_spin.setSuffix(" 厘米")
        margin_layout.addRow("左边距:", self.margin_left_spin)
        
        self.margin_right_spin = QDoubleSpinBox()
        self.margin_right_spin.setRange(0, 10)
        self.margin_right_spin.setValue(3.17)
        self.margin_right_spin.setSuffix(" 厘米")
        margin_layout.addRow("右边距:", self.margin_right_spin)
        
        margin_group.setLayout(margin_layout)
        layout.addWidget(margin_group)
        
        # 页眉页脚
        header_footer_group = QGroupBox("页眉页脚")
        header_footer_layout = QFormLayout()
        
        self.header_font_combo = QFontComboBox()
        self.header_font_combo.setFontFilters(QFontComboBox.AllFonts)
        header_footer_layout.addRow("页眉字体:", self.header_font_combo)
        
        self.header_font_size_spin = QSpinBox()
        self.header_font_size_spin.setRange(8, 48)
        self.header_font_size_spin.setValue(9)
        header_footer_layout.addRow("页眉字体大小:", self.header_font_size_spin)
        
        self.footer_font_combo = QFontComboBox()
        self.footer_font_combo.setFontFilters(QFontComboBox.AllFonts)
        header_footer_layout.addRow("页脚字体:", self.footer_font_combo)
        
        self.footer_font_size_spin = QSpinBox()
        self.footer_font_size_spin.setRange(8, 48)
        self.footer_font_size_spin.setValue(9)
        header_footer_layout.addRow("页脚字体大小:", self.footer_font_size_spin)
        
        self.header_text_edit = QLineEdit()
        self.header_text_edit.setPlaceholderText("页眉文本")
        header_footer_layout.addRow("页眉文本:", self.header_text_edit)
        
        self.footer_text_edit = QLineEdit()
        self.footer_text_edit.setPlaceholderText("页脚文本")
        header_footer_layout.addRow("页脚文本:", self.footer_text_edit)
        
        header_footer_group.setLayout(header_footer_layout)
        layout.addWidget(header_footer_group)
        
        layout.addStretch()
        return widget
    
    def choose_color(self, button):
        """选择颜色"""
        color = QColorDialog.getColor()
        if color.isValid():
            button.setStyleSheet(f"background-color: {color.name()}")
            button.setProperty("color", color.name())
    
    def load_template_preview(self, template_name):
        """加载模板预览"""
        templates = {
            "学术论文模板": """
学术论文模板预览:
- 标题格式: 居中，加粗，16pt
- 作者信息: 居中，12pt
- 摘要: 段落首行缩进2字符，12pt
- 关键词: 段落首行缩进2字符，12pt
- 正文: 段落首行缩进2字符，12pt，1.5倍行距
- 一级标题: 居中，加粗，16pt
- 二级标题: 左对齐，加粗，14pt
- 三级标题: 左对齐，加粗，13pt
- 参考文献: 段落首行缩进2字符，10pt
            """,
            "报告模板": """
报告模板预览:
- 标题格式: 居中，加粗，18pt
- 章节标题: 居中，加粗，16pt
- 小节标题: 左对齐，加粗，14pt
- 正文: 段落首行缩进2字符，12pt，1.5倍行距
- 列表项: 12pt，左对齐
- 图表标题: 居中，12pt，加粗
- 表格内容: 11pt，居中
            """,
            "小说模板": """
小说模板预览:
- 书名: 居中，加粗，20pt
- 作者信息: 居中，14pt
- 章节标题: 居中，16pt
- 正文: 段落首行缩进2字符，12pt，1.5倍行距
- 对话: 与正文相同，可加引号标识
- 注释: 脚注格式，9pt
            """,
            "简历模板": """
简历模板预览:
- 姓名: 居中，加粗，18pt
- 联系方式: 居中，12pt
- 栏目标题: 左对齐，加粗，14pt
- 内容段落: 左对齐，11pt
- 项目列表: 11pt，左对齐，项目符号
- 日期: 右对齐，11pt
            """,
            "信函模板": """
信函模板预览:
- 发件人信息: 左对齐，12pt
- 日期: 右对齐，12pt
- 收件人信息: 左对齐，12pt
- 称呼: 左对齐，12pt
- 正文: 段落首行缩进2字符，12pt，1.5倍行距
- 结尾问候语: 左对齐，12pt
- 署名: 左对齐，12pt
            """,
            "默认模板": """
默认模板预览:
- 标题: 居中，加粗，16pt
- 副标题: 居中，14pt
- 正文: 段落首行缩进2字符，12pt，1.5倍行距
- 一级标题: 左对齐，加粗，16pt
- 二级标题: 左对齐，加粗，14pt
- 三级标题: 左对齐，加粗，13pt
- 列表项: 12pt，左对齐
- 表格: 12pt，居中
- 图片说明: 居中，10pt
            """
        }
        
        # 更新预览文本
        if hasattr(self, 'preview_text'):
            self.preview_text.setPlainText(templates.get(template_name, "模板预览不可用"))
    
    def apply_preset_template(self):
        """应用预设模板"""
        template_name = self.template_combo.currentText()
        
        if template_name == "学术论文模板":
            self.apply_academic_template()
        elif template_name == "报告模板":
            self.apply_report_template()
        elif template_name == "小说模板":
            self.apply_novel_template()
        elif template_name == "简历模板":
            self.apply_resume_template()
        elif template_name == "信函模板":
            self.apply_letter_template()
        elif template_name == "默认模板":
            self.load_default_config()
    
    def apply_academic_template(self):
        """应用学术论文模板"""
        # 基础文本与段落样式
        self.chinese_font_combo.setCurrentFont(QFont("宋体"))
        self.english_font_combo.setCurrentFont(QFont("Times New Roman"))
        self.number_font_combo.setCurrentFont(QFont("Times New Roman"))
        self.first_line_indent_spin.setValue(2.0)
        self.line_spacing_combo.setCurrentText("1.5")
        self.paragraph_spacing_spin.setValue(6)
        self.justify_radio.setChecked(True)
        
        # 标题样式
        for level in range(1, 10):
            if f"标题{level}" in self.heading_configs:
                if level == 1:
                    self.heading_configs[f"标题{level}"]['size'].setValue(16)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("居中对齐")
                elif level == 2:
                    self.heading_configs[f"标题{level}"]['size'].setValue(14)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
                elif level == 3:
                    self.heading_configs[f"标题{level}"]['size'].setValue(13)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
                else:
                    self.heading_configs[f"标题{level}"]['size'].setValue(12)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(False)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
        
        # 其他样式
        self.footnote_font_size_spin.setValue(9)
        self.footnote_italic_check.setChecked(True)
        self.footnote_ref_superscript_check.setChecked(True)
        
    def apply_report_template(self):
        """应用报告模板"""
        # 基础文本与段落样式
        self.chinese_font_combo.setCurrentFont(QFont("宋体"))
        self.english_font_combo.setCurrentFont(QFont("Arial"))
        self.number_font_combo.setCurrentFont(QFont("Arial"))
        self.first_line_indent_spin.setValue(2.0)
        self.line_spacing_combo.setCurrentText("1.5")
        self.paragraph_spacing_spin.setValue(6)
        self.justify_radio.setChecked(True)
        
        # 标题样式
        for level in range(1, 10):
            if f"标题{level}" in self.heading_configs:
                if level == 1:
                    self.heading_configs[f"标题{level}"]['size'].setValue(18)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("居中对齐")
                elif level == 2:
                    self.heading_configs[f"标题{level}"]['size'].setValue(16)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("居中对齐")
                else:
                    self.heading_configs[f"标题{level}"]['size'].setValue(14)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
        
        # 表格样式
        self.table_caption_bold_check.setChecked(True)
        self.table_font_size_spin.setValue(11)
    
    def apply_novel_template(self):
        """应用小说模板"""
        # 基础文本与段落样式
        self.chinese_font_combo.setCurrentFont(QFont("宋体"))
        self.english_font_combo.setCurrentFont(QFont("Times New Roman"))
        self.number_font_combo.setCurrentFont(QFont("Times New Roman"))
        self.first_line_indent_spin.setValue(2.0)
        self.line_spacing_combo.setCurrentText("1.5")
        self.paragraph_spacing_spin.setValue(6)
        self.justify_radio.setChecked(True)
        
        # 标题样式
        for level in range(1, 10):
            if f"标题{level}" in self.heading_configs:
                if level == 1:
                    self.heading_configs[f"标题{level}"]['size'].setValue(20)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("居中对齐")
                elif level == 2:
                    self.heading_configs[f"标题{level}"]['size'].setValue(16)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("居中对齐")
                else:
                    self.heading_configs[f"标题{level}"]['size'].setValue(14)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
    
    def apply_resume_template(self):
        """应用简历模板"""
        # 基础文本与段落样式
        self.chinese_font_combo.setCurrentFont(QFont("宋体"))
        self.english_font_combo.setCurrentFont(QFont("Arial"))
        self.number_font_combo.setCurrentFont(QFont("Arial"))
        self.first_line_indent_spin.setValue(0)  # 简历通常不首行缩进
        self.line_spacing_combo.setCurrentText("1.15")
        self.paragraph_spacing_spin.setValue(3)
        self.left_radio.setChecked(True)  # 左对齐
        
        # 标题样式
        for level in range(1, 10):
            if f"标题{level}" in self.heading_configs:
                if level == 1:
                    self.heading_configs[f"标题{level}"]['size'].setValue(18)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("居中对齐")
                else:
                    self.heading_configs[f"标题{level}"]['size'].setValue(14)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
    
    def apply_letter_template(self):
        """应用信函模板"""
        # 基础文本与段落样式
        self.chinese_font_combo.setCurrentFont(QFont("宋体"))
        self.english_font_combo.setCurrentFont(QFont("Times New Roman"))
        self.number_font_combo.setCurrentFont(QFont("Times New Roman"))
        self.first_line_indent_spin.setValue(2.0)
        self.line_spacing_combo.setCurrentText("1.5")
        self.paragraph_spacing_spin.setValue(6)
        self.left_radio.setChecked(True)  # 左对齐
        
        # 标题样式
        for level in range(1, 10):
            if f"标题{level}" in self.heading_configs:
                if level == 1:
                    self.heading_configs[f"标题{level}"]['size'].setValue(16)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
                else:
                    self.heading_configs[f"标题{level}"]['size'].setValue(14)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
    

    
    
    
    def load_default_config(self):
        """加载默认配置"""
        # 基础文本与段落样式默认值
        self.chinese_font_combo.setCurrentFont(QFont("宋体"))
        self.english_font_combo.setCurrentFont(QFont("Times New Roman"))
        self.number_font_combo.setCurrentFont(QFont("Times New Roman"))
        self.paragraph_font_combo.setCurrentFont(QFont("宋体"))
        self.body_font_combo.setCurrentFont(QFont("宋体"))
        self.justify_radio.setChecked(True)
        
        # 标题样式默认值
        for level in range(1, 10):
            if f"标题{level}" in self.heading_configs:
                if level == 1:
                    self.heading_configs[f"标题{level}"]['size'].setValue(16)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("居中对齐")
                elif level == 2:
                    self.heading_configs[f"标题{level}"]['size'].setValue(14)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
                elif level == 3:
                    self.heading_configs[f"标题{level}"]['size'].setValue(13)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(True)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
                else:
                    self.heading_configs[f"标题{level}"]['size'].setValue(12)
                    self.heading_configs[f"标题{level}"]['bold'].setChecked(False)
                    self.heading_configs[f"标题{level}"]['align'].setCurrentText("左对齐")
        
        # 列表样式默认值
        self.ordered_font_combo.setCurrentFont(QFont("宋体"))
        self.unordered_font_combo.setCurrentFont(QFont("宋体"))
        
        # 引用与交互元素样式默认值
        self.link_font_combo.setCurrentFont(QFont("宋体"))
        self.link_underline_check.setChecked(True)
        self.footnote_font_combo.setCurrentFont(QFont("宋体"))
        self.footnote_italic_check.setChecked(True)
        self.footnote_ref_superscript_check.setChecked(True)
        
        # 非文本与布局样式块默认值
        self.block_font_combo.setCurrentFont(QFont("宋体"))
        self.table_caption_font_combo.setCurrentFont(QFont("宋体"))
        self.table_caption_bold_check.setChecked(True)
        self.table_font_combo.setCurrentFont(QFont("宋体"))
        self.table_border_check.setChecked(True)
        self.figure_caption_font_combo.setCurrentFont(QFont("宋体"))
        self.figure_caption_italic_check.setChecked(True)
        self.figure_caption_position_combo.setCurrentText("图片下方")
        
        # 加载默认模板预览
        self.load_template_preview("默认模板")
    
    
    
    def export_template(self):
        """导出模板文件"""
        # 先收集当前配置
        self.collect_current_config()
        
        # 创建模板文档
        doc = self.create_template_file()
        
        # 选择保存位置
        file_path, _ = QFileDialog.getSaveFileName(
            self, '保存模板文件', '', 'Word文档 (*.docx)'
        )
        
        if file_path:
            try:
                doc.save(file_path)
                QMessageBox.information(self, '成功', f'模板文件已保存到：{file_path}')
            except Exception as e:
                QMessageBox.critical(self, '错误', f'保存模板文件失败：\n{str(e)}')
    
    def collect_current_config(self):
        """收集当前配置"""
        # 收集基础文本与段落样式
        self.config['basic_text'] = {
            'paragraph': {
                'font': self.paragraph_font_combo.currentFont().family(),
                'size': self.paragraph_font_size_spin.value(),
                'bold': self.paragraph_bold_check.isChecked(),
                'italic': self.paragraph_italic_check.isChecked()
            },
            'body': {
                'font': self.body_font_combo.currentFont().family(),
                'size': self.body_font_size_spin.value(),
                'bold': self.body_bold_check.isChecked(),
                'italic': self.body_italic_check.isChecked()
            },
            'char': {
                'chinese_font': self.chinese_font_combo.currentFont().family(),
                'english_font': self.english_font_combo.currentFont().family(),
                'number_font': self.number_font_combo.currentFont().family(),
                'size': self.char_font_size_spin.value()
            },
            'format': {
                'first_line_indent': self.first_line_indent_spin.value(),
                'line_spacing': float(self.line_spacing_combo.currentText()),
                'paragraph_spacing': self.paragraph_spacing_spin.value(),
                'alignment': self.align_group.checkedId()
            }
        }
        
        # 收集标题与层级样式
        self.config['headings'] = {}
        for level in range(1, 10):
            if f"标题{level}" in self.heading_configs:
                controls = self.heading_configs[f"标题{level}"]
                self.config['headings'][f"标题{level}"] = {
                    'font': controls['font'].currentFont().family(),
                    'size': controls['size'].value(),
                    'bold': controls['bold'].isChecked(),
                    'italic': controls['italic'].isChecked(),
                    'underline': controls['underline'].isChecked(),
                    'align': controls['align'].currentText()
                }
        
        # 收集列表样式
        self.config['lists'] = {
            'ordered': {
                'font': self.ordered_font_combo.currentFont().family(),
                'size': self.ordered_font_size_spin.value(),
                'bold': self.ordered_bold_check.isChecked(),
                'italic': self.ordered_italic_check.isChecked(),
                'format': self.ordered_format_combo.currentText()
            },
            'unordered': {
                'font': self.unordered_font_combo.currentFont().family(),
                'size': self.unordered_font_size_spin.value(),
                'bold': self.unordered_bold_check.isChecked(),
                'italic': self.unordered_italic_check.isChecked(),
                'format': self.unordered_format_combo.currentText()
            }
        }
        
        # 收集引用与交互元素样式
        self.config['references'] = {
            'link': {
                'font': self.link_font_combo.currentFont().family(),
                'size': self.link_font_size_spin.value(),
                'bold': self.link_bold_check.isChecked(),
                'italic': self.link_italic_check.isChecked(),
                'underline': self.link_underline_check.isChecked(),
                'color': getattr(self.link_color_button, 'property', lambda x: None)('color') or '#0000FF'
            },
            'footnote': {
                'font': self.footnote_font_combo.currentFont().family(),
                'size': self.footnote_font_size_spin.value(),
                'bold': self.footnote_bold_check.isChecked(),
                'italic': self.footnote_italic_check.isChecked(),
                'underline': self.footnote_underline_check.isChecked()
            },
            'footnote_ref': {
                'size': self.footnote_ref_font_size_spin.value(),
                'superscript': self.footnote_ref_superscript_check.isChecked(),
                'bold': self.footnote_ref_bold_check.isChecked(),
                'italic': self.footnote_ref_italic_check.isChecked()
            }
        }
        
        # 收集非文本与布局样式块
        self.config['layout_elements'] = {
            'block_text': {
                'font': self.block_font_combo.currentFont().family(),
                'size': self.block_font_size_spin.value(),
                'bold': self.block_bold_check.isChecked(),
                'italic': self.block_italic_check.isChecked(),
                'indent': self.block_indent_spin.value()
            },
            'horizontal_line': {
                'thickness': self.hr_thickness_spin.value(),
                'width': self.hr_width_spin.value(),
                'color': getattr(self.hr_color_button, 'property', lambda x: None)('color') or '#000000'
            },
            'table_caption': {
                'font': self.table_caption_font_combo.currentFont().family(),
                'size': self.table_caption_font_size_spin.value(),
                'bold': self.table_caption_bold_check.isChecked(),
                'italic': self.table_caption_italic_check.isChecked(),
                'position': self.table_caption_position_combo.currentText()
            },
            'table': {
                'font': self.table_font_combo.currentFont().family(),
                'size': self.table_font_size_spin.value(),
                'header_bold': self.table_header_bold_check.isChecked(),
                'border': self.table_border_check.isChecked()
            },
            'figure_caption': {
                'font': self.figure_caption_font_combo.currentFont().family(),
                'size': self.figure_caption_font_size_spin.value(),
                'bold': self.figure_caption_bold_check.isChecked(),
                'italic': self.figure_caption_italic_check.isChecked(),
                'position': self.figure_caption_position_combo.currentText()
            }
        }
        
        # 收集页面设置
        self.config['page_settings'] = {
            'size': self.page_size_combo.currentText(),
            'orientation': self.orientation_group.checkedId(),
            'margins': {
                'top': self.margin_top_spin.value(),
                'bottom': self.margin_bottom_spin.value(),
                'left': self.margin_left_spin.value(),
                'right': self.margin_right_spin.value()
            },
            'header_footer': {
                'header_font': self.header_font_combo.currentFont().family(),
                'header_size': self.header_font_size_spin.value(),
                'header_text': self.header_text_edit.text(),
                'footer_font': self.footer_font_combo.currentFont().family(),
                'footer_size': self.footer_font_size_spin.value(),
                'footer_text': self.footer_text_edit.text()
            }
        }
    

    
    def create_template_file(self):
        """根据当前配置创建模板文件"""
        # 创建一个新文档
        doc = Document()
        
        # 应用页面设置
        self.apply_page_settings_to_doc(doc)
        
        # 应用样式设置
        self.apply_style_settings_to_doc(doc)
        
        # 添加预置内容
        self.add_preset_content(doc)
        
        return doc
    
    def apply_page_settings_to_doc(self, doc):
        """将页面设置应用到文档"""
        section = doc.sections[0] if doc.sections else doc.add_section()
        
        # 获取页面设置
        page_settings = self.config.get('page_settings', {})
        
        # 页面尺寸
        page_size = page_settings.get('size', 'A4')
        if page_size == 'A4':
            section.page_width, section.page_height = Inches(8.27), Inches(11.69)
        elif page_size == 'A3':
            section.page_width, section.page_height = Inches(11.69), Inches(16.54)
        elif page_size == 'A5':
            section.page_width, section.page_height = Inches(5.83), Inches(8.27)
        elif page_size == 'Letter':
            section.page_width, section.page_height = Inches(8.5), Inches(11.0)
        
        # 页面方向
        if page_settings.get('orientation', 0) == 1:  # 1表示横向
            section.page_width, section.page_height = section.page_height, section.page_width
        
        # 页边距
        margins = page_settings.get('margins', {})
        section.top_margin = Inches(margins.get('top', 2.54) / 2.54)
        section.bottom_margin = Inches(margins.get('bottom', 2.54) / 2.54)
        section.left_margin = Inches(margins.get('left', 3.17) / 2.54)
        section.right_margin = Inches(margins.get('right', 3.17) / 2.54)
    
    def apply_style_settings_to_doc(self, doc):
        """将样式设置应用到文档"""
        # 应用基础文本与段落样式
        basic_text = self.config.get('basic_text', {})
        
        # 应用正文样式
        normal_style = doc.styles['Normal']
        if 'body' in basic_text:
            body_info = basic_text['body']
            font = normal_style.font
            font.name = body_info.get('font', '宋体')
            font.size = Pt(body_info.get('size', 12))
            font.bold = body_info.get('bold', False)
            font.italic = body_info.get('italic', False)
        
        # 设置中英文字体
        if 'char' in basic_text:
            char_info = basic_text['char']
            normal_style.font.name = char_info.get('chinese_font', '宋体')
            normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), char_info.get('chinese_font', '宋体'))
            normal_style._element.rPr.rFonts.set(qn('w:ascii'), char_info.get('english_font', 'Times New Roman'))
            normal_style._element.rPr.rFonts.set(qn('w:hAnsi'), char_info.get('english_font', 'Times New Roman'))
            normal_style._element.rPr.rFonts.set(qn('w:cs'), char_info.get('number_font', 'Times New Roman'))
        
        # 应用段落格式
        if 'format' in basic_text:
            format_info = basic_text['format']
            paragraph_format = normal_style.paragraph_format
            paragraph_format.first_line_indent = Pt(format_info.get('first_line_indent', 2) * 12)  # 假设12pt为一个字符宽度
            paragraph_format.line_spacing = format_info.get('line_spacing', 1.5)
            paragraph_format.space_before = Pt(format_info.get('paragraph_spacing', 6))
            
            # 设置对齐方式
            alignment = format_info.get('alignment', 3)  # 默认两端对齐
            if alignment == 0:  # 左对齐
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif alignment == 1:  # 居中对齐
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == 2:  # 右对齐
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:  # 两端对齐
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # 应用标题样式
        headings = self.config.get('headings', {})
        for level in range(1, 10):
            style_name = f'标题{level}'
            if style_name in headings:
                try:
                    heading_style = doc.styles[f'Heading {level}']
                except:
                    # 如果没有对应的标题样式，创建一个
                    heading_style = doc.styles.add_style(f'Heading {level}', WD_STYLE_TYPE.PARAGRAPH)
                
                heading_info = headings[style_name]
                font = heading_style.font
                font.name = heading_info.get('font', '宋体')
                font.size = Pt(heading_info.get('size', 12))
                font.bold = heading_info.get('bold', False)
                font.italic = heading_info.get('italic', False)
                
                # 应用对齐方式
                align_text = heading_info.get('align', '左对齐')
                if align_text == '左对齐':
                    heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif align_text == '居中对齐':
                    heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif align_text == '右对齐':
                    heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:  # 两端对齐
                    heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    def add_preset_content(self, doc):
        """添加预置内容，完整展示所有配置的样式项"""
        # 添加文档标题
        title = doc.add_heading('文档标题', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加副标题
        subtitle = doc.add_heading('副标题', level=2)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加作者信息
        author_para = doc.add_paragraph('作者：您的姓名')
        author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加日期
        date_para = doc.add_paragraph('日期：2023年12月31日')
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加空行
        doc.add_paragraph('')
        
        # 添加摘要标题
        doc.add_heading('摘要', level=2)
        
        # 添加摘要内容
        abstract_para = doc.add_paragraph('这是一个示例摘要。摘要应该简明扼要地概括文档的主要内容，通常在200-300字之间。'
                                         '本模板展示了所有可能的文档元素样式，包括各级标题、段落、列表、表格、图片说明、脚注、超链接等。')
        
        # 添加空行
        doc.add_paragraph('')
        
        # 添加关键词标题
        doc.add_heading('关键词', level=3)
        
        # 添加关键词内容
        keywords_para = doc.add_paragraph('关键词1, 关键词2, 关键词3')
        
        # 添加空行
        doc.add_paragraph('')
        
        # 添加正文标题
        doc.add_heading('1. 引言', level=1)
        
        # 添加正文内容
        intro_para = doc.add_paragraph('这是引言部分的内容。引言应该介绍文档的背景、目的和主要内容。'
                                        '在学术写作中，引言通常还会简要介绍论文的结构和研究方法。'
                                        '本段落展示了正文的基本样式，包括字体、字号、行距、首行缩进等格式设置。')
        
        # 添加二级标题
        doc.add_heading('1.1 研究背景', level=2)
        
        # 添加背景内容
        background_para = doc.add_paragraph('这是研究背景的内容。背景部分应该详细说明研究的动机和意义，'
                                           '以及当前领域的研究现状和存在的问题。本段落展示了二级标题下的正文样式。')
        
        # 添加三级标题
        doc.add_heading('1.1.1 相关工作', level=3)
        
        # 添加相关工作内容
        related_work_para = doc.add_paragraph('这是相关工作的内容。相关工作应该综述与本研究相关的已有工作，'
                                            '并指出它们的优势和不足。本段落展示了三级标题下的正文样式。')
        
        # 添加四级标题
        doc.add_heading('1.1.1.1 具体研究方向', level=4)
        
        # 添加四级标题下的内容
        specific_para = doc.add_paragraph('这是四级标题下的内容，展示了更深层次的结构组织方式。')
        
        # 添加五级标题
        doc.add_heading('1.1.1.1.1 研究细节', level=5)
        
        # 添加五级标题下的内容
        detail_para = doc.add_paragraph('这是五级标题下的内容，展示了最精细的文档结构层次。')
        
        # 添加二级标题
        doc.add_heading('1.2 研究目的', level=2)
        
        # 添加研究目的内容
        purpose_para = doc.add_paragraph('这是研究目的的内容。研究目的应该明确说明本研究的目标和预期成果。'
                                        '本段落再次展示了二级标题下的正文样式，保持文档格式的统一性。')
        
        # 添加正文标题
        doc.add_heading('2. 方法', level=1)
        
        # 添加方法内容
        method_para = doc.add_paragraph('这是方法部分的内容。方法部分应该详细描述研究的方法、实验设计和数据分析方法。'
                                       '本段落展示了一级标题下的正文样式。')
        
        # 添加二级标题
        doc.add_heading('2.1 实验设计', level=2)
        
        # 添加实验设计内容
        exp_design_para = doc.add_paragraph('这是实验设计的内容。实验设计应该详细描述实验的设置、参与者、材料和流程。')
        
        # 添加三级标题
        doc.add_heading('2.1.1 实验材料', level=3)
        
        # 添加实验材料内容
        materials_para = doc.add_paragraph('这是实验材料的描述，包括使用的仪器、试剂、软件等。')
        
        # 添加有序列表标题
        doc.add_heading('2.2 实验步骤', level=2)
        
        # 添加有序列表
        doc.add_paragraph('实验步骤如下：', style='List Paragraph')
        doc.add_paragraph('1. 步骤一：准备实验材料', style='List Paragraph')
        doc.add_paragraph('2. 步骤二：设置实验环境', style='List Paragraph')
        doc.add_paragraph('3. 步骤三：进行实验', style='List Paragraph')
        doc.add_paragraph('4. 步骤四：收集和分析数据', style='List Paragraph')
        doc.add_paragraph('5. 步骤五：得出结论', style='List Paragraph')
        
        # 添加无序列表标题
        doc.add_heading('2.3 注意事项', level=2)
        
        # 添加无序列表
        doc.add_paragraph('实验注意事项：', style='List Bullet')
        doc.add_paragraph('• 注意事项一：安全第一', style='List Bullet')
        doc.add_paragraph('• 注意事项二：准确记录数据', style='List Bullet')
        doc.add_paragraph('• 注意事项三：保持实验环境整洁', style='List Bullet')
        doc.add_paragraph('• 注意事项四：遵守实验规程', style='List Bullet')
        
        # 添加正文标题
        doc.add_heading('3. 结果与讨论', level=1)
        
        # 添加结果内容
        result_para = doc.add_paragraph('这是结果与讨论部分的内容。结果部分应该呈现研究的主要发现，'
                                       '而讨论部分应该解释这些发现的意义和可能的原因。')
        
        # 添加块引用
        blockquote = doc.add_paragraph()
        blockquote.add_run('这是一个块引用示例。块引用通常用于引用他人的观点或重要的引用内容。')
        blockquote.paragraph_format.left_indent = Inches(0.5)
        blockquote.paragraph_format.right_indent = Inches(0.5)
        
        # 添加内联格式示例
        format_example = doc.add_paragraph()
        format_example.add_run('这是普通文本，')
        format_example.add_run('这是粗体文本，').bold = True
        format_example.add_run('这是斜体文本，').italic = True
        format_example.add_run('这是下划线文本，').underline = True
        
        # 添加二级标题
        doc.add_heading('3.1 数据分析', level=2)
        
        # 添加表格标题
        doc.add_paragraph('表1：实验数据统计', style='Caption')
        
        # 添加一个示例表格
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '实验组'
        hdr_cells[2].text = '对照组'
        hdr_cells[3].text = '变化率'
        
        # 数据行
        row1_cells = table.rows[1].cells
        row1_cells[0].text = '平均值'
        row1_cells[1].text = '85.6'
        row1_cells[2].text = '72.3'
        row1_cells[3].text = '+18.4%'
        
        row2_cells = table.rows[2].cells
        row2_cells[0].text = '标准差'
        row2_cells[1].text = '5.2'
        row2_cells[2].text = '6.8'
        row2_cells[3].text = '-23.5%'
        
        row3_cells = table.rows[3].cells
        row3_cells[0].text = '样本量'
        row3_cells[1].text = '120'
        row3_cells[2].text = '115'
        row3_cells[3].text = '+4.3%'
        
        # 添加图片说明
        doc.add_paragraph('图1：实验结果对比图', style='Caption')
        
        # 添加二级标题
        doc.add_heading('3.2 讨论', level=2)
        
        # 添加讨论内容，包含脚注引用
        discussion_para = doc.add_paragraph('这是讨论部分的内容。讨论应该解释结果的意义，并与之前的研究进行比较。'
                                          '如参考文献[1]所示，我们的结果与之前的研究基本一致。请参见脚注1获取更多信息。')
        
        # 添加脚注
        footnote_ref = doc.add_paragraph().add_run('1')
        footnote_ref.font.superscript = True
        footnote_ref.font.bold = True
        doc.add_paragraph('1. 这是脚注内容。脚注用于提供额外的信息或引用来源，通常出现在页面的底部。')
        
        # 添加超链接示例
        link_para = doc.add_paragraph()
        link_para.add_run('更多信息请访问：')
        link_run = link_para.add_run('示例链接')
        # 实际应用中，这里应该添加真正的超链接
        link_run.underline = True
        link_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
        
        # 添加水平线
        doc.add_paragraph('_ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _')
        
        # 添加正文标题
        doc.add_heading('4. 结论', level=1)
        
        # 添加结论内容
        conclusion_para = doc.add_paragraph('这是结论部分的内容。结论应该总结研究的主要发现和贡献，'
                                           '并指出研究的局限性和未来研究方向。')
        
        # 添加二级标题
        doc.add_heading('4.1 研究贡献', level=2)
        
        # 添加贡献内容
        contribution_para = doc.add_paragraph('本研究的主要贡献包括：'
                                             '第一，提出了一种新的方法；'
                                             '第二，验证了方法的有效性；'
                                             '第三，指出了未来的研究方向。')
        
        # 添加二级标题
        doc.add_heading('4.2 研究局限性', level=2)
        
        # 添加局限性内容
        limitation_para = doc.add_paragraph('本研究的局限性包括：样本量较小、实验环境单一等。')
        
        # 添加二级标题
        doc.add_heading('4.3 未来工作', level=2)
        
        # 添加未来工作内容
        future_para = doc.add_paragraph('未来的研究方向包括：扩大样本量、改进实验设计、探索更多应用场景等。')
        
        # 添加水平线
        doc.add_paragraph('_ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _')
        
        # 添加正文标题
        doc.add_heading('参考文献', level=1)
        
        # 添加参考文献内容
        ref1 = doc.add_paragraph('[1] 作者. 文献标题. 期刊名称, 年份, 卷(期): 页码.')
        ref2 = doc.add_paragraph('[2] 作者. 书籍名称. 出版社, 年份.')
        ref3 = doc.add_paragraph('[3] 作者. 论文标题. 会议名称, 年份: 页码.')
        ref4 = doc.add_paragraph('[4] 作者. 报告标题. 机构名称, 年份.')
        ref5 = doc.add_paragraph('[5] 作者. 网页标题. 网站名称, 发布日期. [访问日期].')
        
        # 添加附录标题
        doc.add_heading('附录', level=1)
        
        # 添加附录内容
        appendix_para = doc.add_paragraph('这是附录部分的内容。附录通常包含原始数据、代码、补充材料等。')
        
        # 添加附录子标题
        doc.add_heading('附录A: 原始数据', level=2)
        
        # 添加附录内容
        data_para = doc.add_paragraph('这是原始数据表格的说明。')
        
        # 添加附录子标题
        doc.add_heading('附录B: 代码示例', level=2)
        
        # 添加附录内容
        code_para = doc.add_paragraph('这是代码示例的说明。')


class PandocGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Pandoc GUI')
        self.setGeometry(100, 100, 600, 400)
        
        # 获取项目根目录
        self.root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.pandoc_path = os.path.join(self.root_dir, 'pandoc', 'pandoc.exe')
        
        # 支持的文件格式
        self.supported_formats = [
            'markdown', 'docx', 'pdf', 'html', 'epub', 'odt', 
            'txt', 'rst', 'json', 'latex', 'xml', 'pptx'
        ]
        
        self.input_file = None
        self.template_file = None
        
        self.init_ui()
        
    def create_menu_bar(self):
        """创建菜单栏"""
        menubar = self.menuBar()
        
        # 文件菜单
        file_menu = menubar.addMenu('文件')
        
        # 工具菜单
        tools_menu = menubar.addMenu('工具')
        
        # 排版配置菜单项
        format_config_action = tools_menu.addAction('排版配置')
        format_config_action.triggered.connect(self.open_format_config_dialog)
    
    def init_ui(self):
        # 创建菜单栏
        self.create_menu_bar()
        
        # 主窗口部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(20)
        
        # 文件选择部分
        file_layout = QVBoxLayout()
        
        # 输入文件选择
        input_layout = QHBoxLayout()
        self.input_label = QLabel('未选择文件')
        self.input_label.setFixedWidth(300)
        input_button = QPushButton('选择输入文件')
        input_button.clicked.connect(self.select_input_file)
        
        input_layout.addWidget(self.input_label)
        input_layout.addWidget(input_button)
        file_layout.addLayout(input_layout)
        
        # 模板文件选择
        template_layout = QHBoxLayout()
        self.template_label = QLabel('未选择模板文件（可选）')
        self.template_label.setFixedWidth(300)
        template_button = QPushButton('选择模板文件')
        template_button.clicked.connect(self.select_template_file)
        
        template_layout.addWidget(self.template_label)
        template_layout.addWidget(template_button)
        file_layout.addLayout(template_layout)
        
        # 格式选择部分
        format_layout = QHBoxLayout()
        format_layout.addWidget(QLabel('输出格式：'))
        
        self.format_combo = QComboBox()
        self.format_combo.addItems(self.supported_formats)
        self.format_combo.setCurrentText('docx')
        format_layout.addWidget(self.format_combo)
        format_layout.addStretch()
        
        
        
        # 按钮区域
        button_layout = QHBoxLayout()
        
        # 转换按钮
        self.convert_button = QPushButton('转换文件')
        self.convert_button.clicked.connect(self.convert_file)
        self.convert_button.setFixedHeight(40)
        self.convert_button.setEnabled(False)
        
        # 创建模板文件按钮
        self.create_template_button = QPushButton('创建模板文件')
        self.create_template_button.clicked.connect(self.open_format_config_dialog)
        self.create_template_button.setFixedHeight(40)
        
        button_layout.addWidget(self.convert_button)
        button_layout.addWidget(self.create_template_button)
        
        # 状态显示
        self.status_label = QLabel('就绪')
        self.status_label.setAlignment(Qt.AlignCenter)
        
        # 添加所有布局到主布局
        main_layout.addLayout(file_layout)
        main_layout.addLayout(format_layout)
        main_layout.addLayout(button_layout)
        main_layout.addWidget(self.status_label)
        main_layout.addStretch()
    
    def select_input_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, '选择输入文件', '', '所有文件 (*.*)'
        )
        if file_path:
            self.input_file = file_path
            self.input_label.setText(os.path.basename(file_path))
            self.convert_button.setEnabled(True)
            self.status_label.setText(f'已选择输入文件：{os.path.basename(file_path)}')
    
    def select_template_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, '选择模板文件', '', 'Word文档 (*.docx)'
        )
        if file_path:
            self.template_file = file_path
            self.template_label.setText(os.path.basename(file_path))
            self.status_label.setText(f'已选择模板文件：{os.path.basename(file_path)}')
    
    def convert_file(self):
        if not self.input_file:
            QMessageBox.warning(self, '警告', '请先选择输入文件')
            return
        
        output_format = self.format_combo.currentText()
        
        # 生成输出文件名
        input_dir = os.path.dirname(self.input_file)
        input_name = os.path.splitext(os.path.basename(self.input_file))[0]
        output_file = os.path.join(input_dir, f'{input_name}_converted.{output_format}')
        
        # 构建pandoc命令
        cmd = [self.pandoc_path, self.input_file, '-o', output_file]
        
        # 如果有模板文件，添加模板参数
        if self.template_file and output_format == 'docx':
            cmd.extend(['--reference-doc', self.template_file])
        
        try:
            self.status_label.setText('正在转换...')
            QApplication.processEvents()
            
            # 执行pandoc命令
            result = subprocess.run(
                cmd, capture_output=True, text=True, check=True
            )
            
            self.status_label.setText(f'转换成功：{os.path.basename(output_file)}')
            QMessageBox.information(self, '成功', f'文件已转换为：{output_file}')
            
        except subprocess.CalledProcessError as e:
            self.status_label.setText('转换失败')
            QMessageBox.critical(
                self, '错误', 
                f'转换失败：\n{e.stderr if e.stderr else str(e)}'
            )
        except Exception as e:
            self.status_label.setText('转换失败')
            QMessageBox.critical(self, '错误', f'发生错误：\n{str(e)}')
    
    def open_format_config_dialog(self):
        """打开排版配置对话框"""
        dialog = FormatConfigDialog(self)
        dialog.exec_()
    

    

    
    


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = PandocGUI()
    window.show()
    sys.exit(app.exec_())