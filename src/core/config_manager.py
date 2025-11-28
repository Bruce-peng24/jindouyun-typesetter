"""
配置管理模块
处理配置的加载、保存和应用
"""

from PyQt5.QtGui import QFont, QColor
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


class ConfigManager:
    """配置管理器"""
    
    def __init__(self):
        self.config = {}
    
    def collect_config(self, widgets):
        """从各个配置组件收集配置"""
        # 收集基础文本与段落样式
        basic_text = widgets['basic_text'].get_config()
        
        # 收集标题与层级样式
        headings = widgets['heading'].get_config()
        
        # 收集列表样式
        lists = widgets['list'].get_config()
        
        # 收集引用与交互元素样式
        references = widgets['reference'].get_config()
        
        # 收集非文本与布局样式块
        layout_elements = widgets['layout'].get_config()
        
        # 收集页面设置
        page_settings = widgets['page'].get_config()
        
        # 保存到配置字典
        self.config = {
            'basic_text': basic_text,
            'headings': headings,
            'lists': lists,
            'references': references,
            'layout_elements': layout_elements,
            'page_settings': page_settings
        }
        
        return self.config
    
    def load_config_to_widgets(self, config, widgets):
        """加载配置到各个组件"""
        # 加载基础文本与段落样式
        if 'basic_text' in config:
            widgets['basic_text'].load_config(config['basic_text'])
        
        # 加载标题与层级样式
        if 'headings' in config:
            widgets['heading'].load_config(config['headings'])
        
        # 加载列表样式
        if 'lists' in config:
            widgets['list'].load_config(config['lists'])
        
        # 加载引用与交互元素样式
        if 'references' in config:
            widgets['reference'].load_config(config['references'])
        
        # 加载非文本与布局样式块
        if 'layout_elements' in config:
            widgets['layout'].load_config(config['layout_elements'])
        
        # 加载页面设置
        if 'page_settings' in config:
            widgets['page'].load_config(config['page_settings'])
    
    def create_template_file(self):
        """根据当前配置创建模板文件"""
        # 创建一个新文档
        doc = Document()
        
        # 应用页面设置
        self.apply_page_settings_to_doc(doc)
        
        # 应用样式设置
        self.apply_style_settings_to_doc(doc)
        
        # 添加预置内容
        self.add_preset_content(doc)
        
        return doc
    
    def apply_page_settings_to_doc(self, doc):
        """将页面设置应用到文档"""
        section = doc.sections[0] if doc.sections else doc.add_section()
        
        # 获取页面设置
        page_settings = self.config.get('page_settings', {})
        
        # 页面尺寸
        page_size = page_settings.get('size', 'A4')
        if page_size == 'A4':
            section.page_width, section.page_height = Inches(8.27), Inches(11.69)
        elif page_size == 'A3':
            section.page_width, section.page_height = Inches(11.69), Inches(16.54)
        elif page_size == 'A5':
            section.page_width, section.page_height = Inches(5.83), Inches(8.27)
        elif page_size == 'Letter':
            section.page_width, section.page_height = Inches(8.5), Inches(11.0)
        
        # 页面方向
        if page_settings.get('orientation', 0) == 1:  # 1表示横向
            section.page_width, section.page_height = section.page_height, section.page_width
        
        # 页边距
        margins = page_settings.get('margins', {})
        section.top_margin = Inches(margins.get('top', 2.54) / 2.54)
        section.bottom_margin = Inches(margins.get('bottom', 2.54) / 2.54)
        section.left_margin = Inches(margins.get('left', 3.17) / 2.54)
        section.right_margin = Inches(margins.get('right', 3.17) / 2.54)
    
    def apply_style_settings_to_doc(self, doc):
        """将样式设置应用到文档"""
        # 应用基础文本与段落样式
        basic_text = self.config.get('basic_text', {})
        
        # 应用正文样式
        normal_style = doc.styles['Normal']
        if 'body' in basic_text:
            body_info = basic_text['body']
            font = normal_style.font
            font.name = body_info.get('font', '宋体')
            font.size = Pt(body_info.get('size', 12))
            font.bold = body_info.get('bold', False)
            font.italic = body_info.get('italic', False)
        
        # 设置中英文字体
        if 'char' in basic_text:
            char_info = basic_text['char']
            normal_style.font.name = char_info.get('chinese_font', '宋体')
            normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), char_info.get('chinese_font', '宋体'))
            normal_style._element.rPr.rFonts.set(qn('w:ascii'), char_info.get('english_font', 'Times New Roman'))
            normal_style._element.rPr.rFonts.set(qn('w:hAnsi'), char_info.get('english_font', 'Times New Roman'))
            normal_style._element.rPr.rFonts.set(qn('w:cs'), char_info.get('number_font', 'Times New Roman'))
        
        # 应用段落格式
        if 'format' in basic_text:
            format_info = basic_text['format']
            paragraph_format = normal_style.paragraph_format
            paragraph_format.first_line_indent = Pt(format_info.get('first_line_indent', 2) * 12)  # 假设12pt为一个字符宽度
            paragraph_format.line_spacing = format_info.get('line_spacing', 1.5)
            paragraph_format.space_before = Pt(format_info.get('paragraph_spacing', 6))
            
            # 设置对齐方式
            alignment = format_info.get('alignment', 3)  # 默认两端对齐
            if alignment == 0:  # 左对齐
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif alignment == 1:  # 居中对齐
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == 2:  # 右对齐
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:  # 两端对齐
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # 应用标题样式
        headings = self.config.get('headings', {})
        for level in range(1, 10):
            style_name = f'标题{level}'
            if style_name in headings:
                try:
                    heading_style = doc.styles[f'Heading {level}']
                except:
                    # 如果没有对应的标题样式，创建一个
                    heading_style = doc.styles.add_style(f'Heading {level}', WD_STYLE_TYPE.PARAGRAPH)
                
                heading_info = headings[style_name]
                font = heading_style.font
                font.name = heading_info.get('font', '宋体')
                font.size = Pt(heading_info.get('size', 12))
                font.bold = heading_info.get('bold', False)
                font.italic = heading_info.get('italic', False)
                
                # 应用对齐方式
                align_text = heading_info.get('align', '左对齐')
                if align_text == '左对齐':
                    heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif align_text == '居中对齐':
                    heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif align_text == '右对齐':
                    heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:  # 两端对齐
                    heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    def add_preset_content(self, doc):
        """添加预置内容，完整展示所有配置的样式项"""
        # 添加文档标题
        title = doc.add_heading('文档标题', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加副标题
        subtitle = doc.add_heading('副标题', level=2)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加作者信息
        author_para = doc.add_paragraph('作者：您的姓名')
        author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加日期
        date_para = doc.add_paragraph('日期：2023年12月31日')
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加空行
        doc.add_paragraph('')
        
        # 添加摘要标题
        doc.add_heading('摘要', level=2)
        
        # 添加摘要内容
        abstract_para = doc.add_paragraph('这是一个示例摘要。摘要应该简明扼要地概括文档的主要内容，通常在200-300字之间。'
                                         '本模板展示了所有可能的文档元素样式，包括各级标题、段落、列表、表格、图片说明、脚注、超链接等。')
        
        # 添加空行
        doc.add_paragraph('')
        
        # 添加关键词标题
        doc.add_heading('关键词', level=3)
        
        # 添加关键词内容
        keywords_para = doc.add_paragraph('关键词1, 关键词2, 关键词3')
        
        # 添加空行
        doc.add_paragraph('')
        
        # 添加正文标题
        doc.add_heading('1. 引言', level=1)
        
        # 添加正文内容
        intro_para = doc.add_paragraph('这是引言部分的内容。引言应该介绍文档的背景、目的和主要内容。'
                                        '在学术写作中，引言通常还会简要介绍论文的结构和研究方法。'
                                        '本段落展示了正文的基本样式，包括字体、字号、行距、首行缩进等格式设置。')
        
        # 添加二级标题
        doc.add_heading('1.1 研究背景', level=2)
        
        # 添加背景内容
        background_para = doc.add_paragraph('这是研究背景的内容。背景部分应该详细说明研究的动机和意义，'
                                           '以及当前领域的研究现状和存在的问题。本段落展示了二级标题下的正文样式。')
        
        # 添加三级标题
        doc.add_heading('1.1.1 相关工作', level=3)
        
        # 添加相关工作内容
        related_work_para = doc.add_paragraph('这是相关工作的内容。相关工作应该综述与本研究相关的已有工作，'
                                            '并指出它们的优势和不足。本段落展示了三级标题下的正文样式。')
        
        # 添加四级标题
        doc.add_heading('1.1.1.1 具体研究方向', level=4)
        
        # 添加四级标题下的内容
        specific_para = doc.add_paragraph('这是四级标题下的内容，展示了更深层次的结构组织方式。')
        
        # 添加五级标题
        doc.add_heading('1.1.1.1.1 研究细节', level=5)
        
        # 添加五级标题下的内容
        detail_para = doc.add_paragraph('这是五级标题下的内容，展示了最精细的文档结构层次。')
        
        # 添加二级标题
        doc.add_heading('1.2 研究目的', level=2)
        
        # 添加研究目的内容
        purpose_para = doc.add_paragraph('这是研究目的的内容。研究目的应该明确说明本研究的目标和预期成果。'
                                        '本段落再次展示了二级标题下的正文样式，保持文档格式的统一性。')
        
        # 添加正文标题
        doc.add_heading('2. 方法', level=1)
        
        # 添加方法内容
        method_para = doc.add_paragraph('这是方法部分的内容。方法部分应该详细描述研究的方法、实验设计和数据分析方法。'
                                       '本段落展示了一级标题下的正文样式。')
        
        # 添加二级标题
        doc.add_heading('2.1 实验设计', level=2)
        
        # 添加实验设计内容
        exp_design_para = doc.add_paragraph('这是实验设计的内容。实验设计应该详细描述实验的设置、参与者、材料和流程。')
        
        # 添加三级标题
        doc.add_heading('2.1.1 实验材料', level=3)
        
        # 添加实验材料内容
        materials_para = doc.add_paragraph('这是实验材料的描述，包括使用的仪器、试剂、软件等。')
        
        # 添加有序列表标题
        doc.add_heading('2.2 实验步骤', level=2)
        
        # 添加有序列表
        doc.add_paragraph('实验步骤如下：', style='List Paragraph')
        doc.add_paragraph('1. 步骤一：准备实验材料', style='List Paragraph')
        doc.add_paragraph('2. 步骤二：设置实验环境', style='List Paragraph')
        doc.add_paragraph('3. 步骤三：进行实验', style='List Paragraph')
        doc.add_paragraph('4. 步骤四：收集和分析数据', style='List Paragraph')
        doc.add_paragraph('5. 步骤五：得出结论', style='List Paragraph')
        
        # 添加无序列表标题
        doc.add_heading('2.3 注意事项', level=2)
        
        # 添加无序列表
        doc.add_paragraph('实验注意事项：', style='List Bullet')
        doc.add_paragraph('• 注意事项一：安全第一', style='List Bullet')
        doc.add_paragraph('• 注意事项二：准确记录数据', style='List Bullet')
        doc.add_paragraph('• 注意事项三：保持实验环境整洁', style='List Bullet')
        doc.add_paragraph('• 注意事项四：遵守实验规程', style='List Bullet')
        
        # 添加正文标题
        doc.add_heading('3. 结果与讨论', level=1)
        
        # 添加结果内容
        result_para = doc.add_paragraph('这是结果与讨论部分的内容。结果部分应该呈现研究的主要发现，'
                                       '而讨论部分应该解释这些发现的意义和可能的原因。')
        
        # 添加块引用
        blockquote = doc.add_paragraph()
        blockquote.add_run('这是一个块引用示例。块引用通常用于引用他人的观点或重要的引用内容。')
        blockquote.paragraph_format.left_indent = Inches(0.5)
        blockquote.paragraph_format.right_indent = Inches(0.5)
        
        # 添加内联格式示例
        format_example = doc.add_paragraph()
        format_example.add_run('这是普通文本，')
        format_example.add_run('这是粗体文本，').bold = True
        format_example.add_run('这是斜体文本，').italic = True
        format_example.add_run('这是下划线文本，').underline = True
        
        # 添加二级标题
        doc.add_heading('3.1 数据分析', level=2)
        
        # 添加表格标题
        doc.add_paragraph('表1：实验数据统计', style='Caption')
        
        # 添加一个示例表格
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '实验组'
        hdr_cells[2].text = '对照组'
        hdr_cells[3].text = '变化率'
        
        # 数据行
        row1_cells = table.rows[1].cells
        row1_cells[0].text = '平均值'
        row1_cells[1].text = '85.6'
        row1_cells[2].text = '72.3'
        row1_cells[3].text = '+18.4%'
        
        row2_cells = table.rows[2].cells
        row2_cells[0].text = '标准差'
        row2_cells[1].text = '5.2'
        row2_cells[2].text = '6.8'
        row2_cells[3].text = '-23.5%'
        
        row3_cells = table.rows[3].cells
        row3_cells[0].text = '样本量'
        row3_cells[1].text = '120'
        row3_cells[2].text = '115'
        row3_cells[3].text = '+4.3%'
        
        # 添加图片说明
        doc.add_paragraph('图1：实验结果对比图', style='Caption')
        
        # 添加二级标题
        doc.add_heading('3.2 讨论', level=2)
        
        # 添加讨论内容，包含脚注引用
        discussion_para = doc.add_paragraph('这是讨论部分的内容。讨论应该解释结果的意义，并与之前的研究进行比较。'
                                          '如参考文献[1]所示，我们的结果与之前的研究基本一致。请参见脚注1获取更多信息。')
        
        # 添加脚注
        footnote_ref = doc.add_paragraph().add_run('1')
        footnote_ref.font.superscript = True
        footnote_ref.font.bold = True
        doc.add_paragraph('1. 这是脚注内容。脚注用于提供额外的信息或引用来源，通常出现在页面的底部。')
        
        # 添加超链接示例
        link_para = doc.add_paragraph()
        link_para.add_run('更多信息请访问：')
        link_run = link_para.add_run('示例链接')
        # 实际应用中，这里应该添加真正的超链接
        link_run.underline = True
        link_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
        
        # 添加水平线
        doc.add_paragraph('_ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _')
        
        # 添加正文标题
        doc.add_heading('4. 结论', level=1)
        
        # 添加结论内容
        conclusion_para = doc.add_paragraph('这是结论部分的内容。结论应该总结研究的主要发现和贡献，'
                                           '并指出研究的局限性和未来研究方向。')
        
        # 添加二级标题
        doc.add_heading('4.1 研究贡献', level=2)
        
        # 添加贡献内容
        contribution_para = doc.add_paragraph('本研究的主要贡献包括：'
                                             '第一，提出了一种新的方法；'
                                             '第二，验证了方法的有效性；'
                                             '第三，指出了未来的研究方向。')
        
        # 添加二级标题
        doc.add_heading('4.2 研究局限性', level=2)
        
        # 添加局限性内容
        limitation_para = doc.add_paragraph('本研究的局限性包括：样本量较小、实验环境单一等。')
        
        # 添加二级标题
        doc.add_heading('4.3 未来工作', level=2)
        
        # 添加未来工作内容
        future_para = doc.add_paragraph('未来的研究方向包括：扩大样本量、改进实验设计、探索更多应用场景等。')
        
        # 添加水平线
        doc.add_paragraph('_ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _')
        
        # 添加正文标题
        doc.add_heading('参考文献', level=1)
        
        # 添加参考文献内容
        ref1 = doc.add_paragraph('[1] 作者. 文献标题. 期刊名称, 年份, 卷(期): 页码.')
        ref2 = doc.add_paragraph('[2] 作者. 书籍名称. 出版社, 年份.')
        ref3 = doc.add_paragraph('[3] 作者. 论文标题. 会议名称, 年份: 页码.')
        ref4 = doc.add_paragraph('[4] 作者. 报告标题. 机构名称, 年份.')
        ref5 = doc.add_paragraph('[5] 作者. 网页标题. 网站名称, 发布日期. [访问日期].')
        
        # 添加附录标题
        doc.add_heading('附录', level=1)
        
        # 添加附录内容
        appendix_para = doc.add_paragraph('这是附录部分的内容。附录通常包含原始数据、代码、补充材料等。')
        
        # 添加附录子标题
        doc.add_heading('附录A: 原始数据', level=2)
        
        # 添加附录内容
        data_para = doc.add_paragraph('这是原始数据表格的说明。')
        
        # 添加附录子标题
        doc.add_heading('附录B: 代码示例', level=2)
        
        # 添加附录内容
        code_para = doc.add_paragraph('这是代码示例的说明。')